from __future__ import annotations

from io import BytesIO

from docx import Document as WordDocument
from fastapi.testclient import TestClient

from harness_app.access.api_gateway.app import create_app
from harness_app.knowledge.chunker import MarkdownChunker
from harness_app.knowledge.models import KnowledgeDocument


def test_markdown_chunker_preserves_titles_and_sources():
    chunker = MarkdownChunker(max_chars=80, overlap_chars=10)
    documents = [
        KnowledgeDocument(
            source_path="docs/guide.md",
            title="guide",
            content="# Intro\n\n## GLM-5\nThis section explains GLM-5 support.\n\n### Frontend\nThe Vue frontend calls the API.",
        )
    ]

    chunks = chunker.chunk_documents(documents)

    assert chunks
    assert chunks[0].source_path == "docs/guide.md"
    assert any(chunk.title == "GLM-5" for chunk in chunks)
    assert all(chunk.chunk_id.startswith("docs/guide.md#") for chunk in chunks)


def test_knowledge_search_endpoint_returns_grounded_sources(monkeypatch):
    monkeypatch.setenv("MYPORTAL_ENV_FILE", "missing.env")
    monkeypatch.delenv("MYPORTAL_LLM_API_KEY", raising=False)
    monkeypatch.delenv("ZHIPUAI_API_KEY", raising=False)
    client = TestClient(create_app())

    response = client.post("/knowledge/search", json={"query": "GLM-5", "top_k": 2})

    assert response.status_code == 200
    payload = response.json()
    assert payload["total_hits"] >= 1
    assert payload["indexed_chunk_count"] >= payload["total_hits"]
    assert payload["sources"][0]["source_path"] in {
        "README.md",
        "MYPORTAL_FUNCTION_ANALYSIS.zh-CN.md",
        "portal-front/README.md",
    }


def test_agent_respond_uses_faiss_grounding(monkeypatch):
    monkeypatch.setenv("MYPORTAL_ENV_FILE", "missing.env")
    monkeypatch.delenv("MYPORTAL_LLM_API_KEY", raising=False)
    monkeypatch.delenv("ZHIPUAI_API_KEY", raising=False)
    client = TestClient(create_app())

    session_response = client.post("/sessions", json={"user_id": "rag-user", "metadata": {}})
    session_id = session_response.json()["session_id"]

    response = client.post(
        "/agent/respond",
        json={
            "user_id": "rag-user",
            "session_id": session_id,
            "message": "search GLM-5 support",
            "channel": "http",
            "metadata": {"knowledge_top_k": 2},
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["tool_results"]
    assert "[faiss_search]" in payload["tool_results"][0]
    assert "source=" in payload["tool_results"][0]


def test_knowledge_upload_list_and_search(monkeypatch, tmp_path):
    monkeypatch.setenv("MYPORTAL_ENV_FILE", "missing.env")
    monkeypatch.setenv("MYPORTAL_KNOWLEDGE_UPLOADS_DIR", str(tmp_path / "uploads"))
    monkeypatch.delenv("MYPORTAL_LLM_API_KEY", raising=False)
    monkeypatch.delenv("ZHIPUAI_API_KEY", raising=False)
    client = TestClient(create_app())

    buffer = BytesIO()
    document = WordDocument()
    document.add_paragraph("OpenHarness knowledge upload supports source-grounded search.")
    document.save(buffer)
    buffer.seek(0)

    upload_response = client.post(
        "/knowledge/upload",
        files=[("files", ("knowledge.docx", buffer.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"))],
    )
    assert upload_response.status_code == 200
    upload_payload = upload_response.json()
    assert upload_payload["files"][0]["status"] == "indexed"
    assert upload_payload["files"][0]["extracted_text_length"] > 0

    list_response = client.get("/knowledge/files")
    assert list_response.status_code == 200
    files_payload = list_response.json()
    assert len(files_payload["files"]) == 1

    reindex_response = client.post("/knowledge/reindex")
    assert reindex_response.status_code == 200
    assert reindex_response.json()["uploaded_file_count"] == 1

    search_response = client.post("/knowledge/search", json={"query": "source-grounded search", "top_k": 3})
    assert search_response.status_code == 200
    search_payload = search_response.json()
    assert search_payload["total_hits"] >= 1
    assert any("source-grounded search" in item["content"] for item in search_payload["sources"])
